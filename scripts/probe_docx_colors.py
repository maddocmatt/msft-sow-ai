"""Probe a docx for the most common font colors so we can detect pink guidance."""

from __future__ import annotations

import sys
from collections import Counter

import docx


def main(path: str) -> None:
    doc = docx.Document(path)
    colors: Counter[str] = Counter()
    samples: dict[str, str] = {}
    for para in doc.paragraphs:
        for run in para.runs:
            c = run.font.color
            if c is None or c.rgb is None:
                continue
            hx = str(c.rgb)
            colors[hx] += 1
            samples.setdefault(hx, (run.text or "").strip()[:90])
    # Walk tables too — Microsoft templates use lots of them
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        c = run.font.color
                        if c is None or c.rgb is None:
                            continue
                        hx = str(c.rgb)
                        colors[hx] += 1
                        samples.setdefault(hx, (run.text or "").strip()[:90])
    for hx, n in colors.most_common(20):
        print(f"{hx}  n={n:4}  '{samples[hx]}'")


if __name__ == "__main__":
    main(sys.argv[1])
